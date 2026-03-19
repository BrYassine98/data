import pandas as pd
from docx import Document
import os

# Load the Jupyter Notebook and convert it to a DataFrame
notebook_file = 'attrition_crispdm.ipynb'

# Function to extract data from notebook and generate the report

def generate_report():
    # Load data from the Jupyter notebook
    # ... (logic to read notebook and process data)

    df = pd.DataFrame(...)  # replace with actual data loading logic

    # Create a Word document
    doc = Document()

    # Add headings and key results to the document
    doc.add_heading('CRISP-DM Attrition Report', level=1)

    # Summary of the DataFrame
    doc.add_heading('Key Results', level=2)
    doc.add_paragraph(f'Shape: {df.shape}')
    doc.add_paragraph(f'Dtypes Summary: {df.dtypes.to_string()}')
    doc.add_paragraph(f'Missing Values: {df.isnull().sum().to_string()}')
    doc.add_paragraph(f'Duplicates: {df.duplicated().sum()}')
    doc.add_paragraph(f'Constant Columns: {[col for col in df.columns if df[col].nunique() == 1]}')

    # Include attrition distribution and baseline accuracy
    doc.add_paragraph('Attrition Distribution: ...')
    doc.add_paragraph('Baseline Accuracy: ...')

    # Embed images if present
    for i in range(1, 11):
        img_path = f'{i}.png'
        if os.path.exists(img_path):
            doc.add_picture(img_path)

    # Save the document
    doc.save('Rapport_CRISPDM_Attrition.docx')

# Call the function to generate the report
generate_report()